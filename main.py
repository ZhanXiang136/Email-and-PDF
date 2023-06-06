from docx import Document
from docx2pdf import convert
import smtplib, ssl
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import os
from dotenv import load_dotenv
import pandas as pd
from os.path import exists
import logging
from datetime import datetime

logger = logging.basicConfig(filename='logs.log', filemode='w', level='DEBUG', format="%(asctime)s:%(levelname)s:%(message)s", datefmt="%m-%d-%y %H:%M:%S %p",)

load_dotenv()
EMAIL = os.getenv('EMAIL')
PASSWORD = os.getenv('PASSWORD')

month = {	'01':'Janauary',
		'02':'February',
		'03':'March',
		'04':'April',
		'05':'May',
		'06':'June',
		'07':'July',
		'08':'August',
		'09':'September',
		'10':'October',
		'11':'November',
		'12':'December'		}

now = datetime.now()
date = f'{month[now.strftime("%m")]} {now.strftime("%d")}, {now.strftime("%Y")}'

def send_mail(sender, reciever, subject, body, password):
    message = MIMEMultipart()
    message["From"] = sender
    message["To"] = reciever
    message["Subject"] = subject

    message.attach(MIMEText(body, "html"))

    filename = "Volunteering_Hours.pdf"

    with open(filename, "rb") as attachment:
        part = MIMEBase("application", "octet-stream")
        part.set_payload(attachment.read())

    encoders.encode_base64(part)
    
    part.add_header("Content-Disposition", f"attachment; filename= {filename}")

    message.attach(part)
    text = message.as_string()

    context = ssl.create_default_context()
    with smtplib.SMTP_SSL("smtp.gmail.com", 465, context=context) as server:
        server.login(sender, password)
        server.sendmail(sender, reciever, text)

def edit_doc(first_name, last_name, hours):
    document = Document("HOPE-Volunteering-Letter.docx")
    paragraph2 = document.paragraphs[2]
    paragraph2.text = paragraph2.text.replace('<<Date>>', date)
    paragraph6 = document.paragraphs[6]
    paragraph6.text = paragraph6.text.replace('<<First Name>>', str(first_name).strip())
    paragraph6.text = paragraph6.text.replace('<<Last Name>>', str(last_name).strip())
    paragraph6.text = paragraph6.text.replace('<<Hours>>', str(hours).strip())
    document.save("Volunteering_Hours.docx")
    convert("Volunteering_Hours.docx") #convert edited docx to pdf

def check_file_existence(text):
    while True:
        user_input = input(f"{text}")
        if exists(user_input):
            return user_input
        else:
            print("Invalid file path")


def determine_correct_pandas_conversion(file):
    extension = file.split('.')[1]
    if extension == 'xlsx':
        return pd.read_excel(file, sheet_name=0)
    elif extension == 'csv':
        return pd.read_csv(file)
    elif extension == 'tsv':
        return pd.read_csv(file, delimiter='\t')
    else:
        print(f'\nError: Invalid extension {extension}')


def main():
    print("Welcome to the Program")
    main_file = 'main.xlsx'
    if not exists('main.xlsx'):
        main_file = check_file_existence('\nEnter the main spreadsheet: ')

    subject = '2022-2023 Hope Volunteering Letter'

    df = determine_correct_pandas_conversion(main_file)
    
    for index, row in df.iterrows():
        print(f"Sending {index} {row['First Name']}")

        with open('body.html', 'r') as file:
            body = file.readlines()
            body = ''.join(body)
            body = body.replace('${First Name}', row['First Name'])

        edit_doc(row['First Name'], row['Last Name'], row['Total Hours'])
        send_mail(EMAIL, row['Email'], subject, body, password=PASSWORD)
        logging.info(f"{row['First Name']} Email sent")
    print('Done')

if __name__ == "__main__":
    main()